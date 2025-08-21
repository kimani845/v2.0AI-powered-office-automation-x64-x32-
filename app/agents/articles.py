import os
import re
import json
from typing import Optional
from datetime import datetime
from pydantic import BaseModel, Field
# Imports for creating .docx files
from docx import Document
from docx.shared import Pt

from app.agents.llm_client import LLMClient

# Pydantic Model 
class ArticleOutput(BaseModel):
    """Defines the structured output for a generated article."""
    title: str = Field(..., description="The compelling title of the article.")
    content: str = Field(..., description="The full, well-structured content of the article, formatted with paragraphs.")

# Helper Functions 
def _clean_json_response(text: str) -> str:
    """Helper to strip markdown backticks or other text from a JSON string."""
    match = re.search(r'```(?:json)?\s*({.*})\s*```', text, re.DOTALL)
    if match:
        return match.group(1).strip()
    return text.strip()

def build_article_prompt(topic: str, length: str = "medium", style: str = "blog post", audience: str = "the general public") -> str:
    """Builds a detailed prompt that instructs the LLM to return a JSON object."""
    return f"""
You are an expert content writer. Your task is to write a high-quality article on the given topic.
You must return your response as a single, valid JSON object, with no other text or markdown before or after it.

The JSON object must have the following exact structure:
{{
    "title": "A compelling and relevant title for the article",
    "content": "The full article content, with proper paragraphs separated by \\n\\n."
}}

---
TOPIC: "{topic}"
STYLE: "{style}"
TARGET AUDIENCE: "{audience}"
DESIRED LENGTH: "{length}" (e.g., a few paragraphs for 'short', a comprehensive piece for 'long')
---

Now, generate the final JSON response based on the instructions above.
"""

# --- NEW: Function to Save Article to .docx ---
def save_article_to_docx(article: ArticleOutput, output_dir: str = "generated_articles") -> str:
    """
    Saves the generated article to a .docx file with basic formatting.

    Args:
        article (ArticleOutput): The structured article data.
        output_dir (str): The directory to save the file in.

    Returns:
        str: The full path to the saved .docx file.
    """
    try:
        # Create the output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)

        # Sanitize the title to create a valid filename
        sanitized_title = re.sub(r'[\\/*?:"<>|]', "", article.title)[:50] # Remove invalid chars and shorten
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{sanitized_title}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)

        # Create a new Word document
        document = Document()

        # Set default font
        style = document.styles['Normal'].font
        style.name = 'Calibri'
        style.size = Pt(12)

        # Add the title as a main heading
        document.add_heading(article.title, level=1)

        # Split the content by double newlines and add as paragraphs
        # This preserves the paragraph structure from the AI's output
        for paragraph_text in article.content.strip().split('\n\n'):
            document.add_paragraph(paragraph_text.strip())

        # Save the document
        document.save(filepath)
        print(f"-> Article successfully saved to: {filepath}")
        return filepath

    except Exception as e:
        print(f"Error saving document: {e}")
        raise

# Class for article generation
class ArticleAgent:
    def __init__(self, llm_client: LLMClient):
        """Initializes the agent with a pre-configured LLMClient."""
        self.llm_client = llm_client
        print("-> Article Agent activated.")

    def create_article(self, topic: str, length: Optional[str] = "medium", style: Optional[str] = "blog post", audience: Optional[str] = "the general public") -> str:
        """
        Generates an article, saves it to a .docx file, and returns the filepath.
        """
        if not topic:
            raise ValueError("A topic must be provided to generate an article.")

        prompt = build_article_prompt(topic, length, style, audience)
        print(f"-> Generating article on '{topic}'...")
        raw_response = self.llm_client.generate_response(prompt, json_mode=True)
        cleaned_response = _clean_json_response(raw_response)
        
        try:
            parsed_json = json.loads(cleaned_response)
            article_output = ArticleOutput.model_validate(parsed_json)
        except (json.JSONDecodeError, TypeError) as e:
            print(f"Error: Failed to parse the LLM's response. Saving an error report.")
            # Create a structured error message to be saved in the docx
            article_output = ArticleOutput(
                title=f"Error Generating Article on '{topic}'", 
                content=(
                    f"The AI's response could not be parsed correctly and an article could not be generated.\n\n"
                    f"Error Details: {e}\n\n"
                    f"Raw AI Response:\n{raw_response}"
                )
            )
        
        # Save the resulting content (either the article or the error report) to a .docx file
        filepath = save_article_to_docx(article_output)
        return filepath