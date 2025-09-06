"""
This script is the backend server for the AI Office Automation add-in.
It uses FastAPI to expose endpoints that perform heavy AI and data processing tasks.
This server is intended to be run as a standalone 64-bit executable.
"""
import os
import sys
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
import docx
import uuid # For generating unique filenames

# # This block makes the script "bundle-aware" and correctly loads the .env file.
# from dotenv import load_dotenv

# if getattr(sys, 'frozen', False):
#     # We are running in a bundled executable.
#     # sys.executable is the path to the .exe file.
#     # We want to find the .env file in the same directory as the .exe.
#     application_path = os.path.dirname(sys.executable)
#     dotenv_path = os.path.join(application_path, '.env')
#     if os.path.exists(dotenv_path):
#         load_dotenv(dotenv_path=dotenv_path)
#     else:
#         # This is a critical error if the .env file is missing in the bundle.
#         # You can add file logging here if you want to debug this on a client's machine.
#         print("FATAL: Bundled .env file not found!")
# else:
#     # We are running in a normal Python environment (development mode).
#     # load_dotenv() will automatically find the .env in the project root.
#     load_dotenv()

# # Path Setup
# if getattr(sys, 'frozen', False):
#     base_path = sys._MEIPASS
# else:
#     base_path = os.path.dirname(os.path.abspath(__file__))

# sys.path.insert(0, base_path)

from dotenv import load_dotenv

def get_base_path():
    """ Get the base path for the application, handling PyInstaller's _MEIPASS folder. """
    if getattr(sys, 'frozen', False):
        # We are running in a bundled executable (_MEIPASS).
        return sys._MEIPASS
    else:
        # We are running in a normal Python environment.
        # The project root is one level up from the wps_addin folder.
        return os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# --- CRITICAL PATHING LOGIC ---
BASE_PATH = get_base_path()

# Load the .env file from the correct base path
dotenv_path = os.path.join(BASE_PATH, '.env')
if os.path.exists(dotenv_path):
    load_dotenv(dotenv_path=dotenv_path)
    print(f"Successfully loaded .env file from: {dotenv_path}")
else:
    print(f"WARNING: .env file not found at {dotenv_path}. Using environment variables.")
    # This allows the script to still work if API keys are set as system environment variables.

# The 'app' directory to the Python path so imports like 'from app.agents...' work
app_dir_path = os.path.join(BASE_PATH, 'app')
if os.path.exists(app_dir_path):
    sys.path.insert(0, app_dir_path)
    print(f"Successfully added '{app_dir_path}' to sys.path.")
else:
    print(f"FATAL: Could not find the 'app' directory at '{app_dir_path}'. Imports will fail.")
    sys.exit(1)
    
# Agent Initializations
try:
    from app.agents.llm_client import LLMClient
    from app.agents.analyzer import StructuredDataAgent
    from app.agents.reports import ReportAgent
    from app.agents.articles import ArticleAgent
    from app.agents.documents import DocumentGenerationAgent, DocumentRequest # DocumentRequest is crucial
except ImportError as e:
    print(f"FATAL: Could not import agent modules. Ensure the 'app' folder is in the same directory. Error: {e}")
    sys.exit(1)

print("Backend Server: Initializing AI agents...")
try:
    llm_client = LLMClient(provider="deepseek")
    report_agent = ReportAgent(llm_client)
    article_agent = ArticleAgent(llm_client)
    data_agent = StructuredDataAgent(llm_client)
    document_agent = DocumentGenerationAgent(llm_client)
    print("Backend Server: AI agents initialized successfully.")
except Exception as e:
    print(f"FATAL: Failed to initialize AI agents. Check API keys in config.json. Error: {e}")
    sys.exit(1)

# Initialize FastAPI Server
app = FastAPI(title="AI Office Automation Backend Server")

# Define Pydantic Models for API Request Bodies
class ProcessRequest(BaseModel):
    prompt: str
    content: str = "" # Optional content field for analysis/summarization

class GeneralResponse(BaseModel):
    result: str

# FastAPI Endpoints
@app.get("/")
def root():
    return {"message": "AI Office Backend Server is running."}

# Dedicated endpoint for Document Generation
# @app.post("/generate_document", response_model=GeneralResponse)
# def generate_document_endpoint(request: DocumentRequest):
#     """Generates a document based on a complete DocumentRequest object."""
#     print("Backend: Received a complete DocumentRequest for generation.")
#     try:
#         # The agent now receives a validated DocumentRequest object directly
#         output_document_obj = document_agent.generate_document(request)
        
#         # document_agent.generate_document returns a docx.Document object.
#         # We need to save it to a temporary file and read its text content.
#         # This is important for sending the text back via FastAPI.
#         temp_file_path = "temp_generated_document.docx"
#         output_document_obj.save(temp_file_path) # Save the docx.Document object to a temporary file
        
#         doc = docx.Document(temp_file_path)
#         full_text = "\n".join([para.text for para in doc.paragraphs])
#         os.remove(temp_file_path) # Clean up the temporary file

#         if not full_text:
#             raise HTTPException(status_code=500, detail="Failed to extract text from generated document.")
#         return GeneralResponse(result=full_text)
#     except Exception as e:
#         print(f"Error in generate_document_endpoint: {e}")
#         raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")


# In backend_server.py, replace the existing /generate_document endpoint
# and add the new dedicated endpoints below it.


# Endpoint to serve downloadable files
@app.get("/download/{filename}")
async def download_file(filename: str):
    """Serves a generated file for download."""
    file_path = os.path.join("generated_documents", filename) # Store in a subfolder
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    # Determine media type based on file extension
    media_type = "application/octet-stream"
    if filename.endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif filename.endswith(".pdf"):
        media_type = "application/pdf"
    elif filename.endswith(".xlsx"):
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    # Add more types as needed

    with open(file_path, "rb") as file:
        content = file.read()
    
    # Optional: Clean up the file after it's served if it's meant to be single-use
    # For now, let's keep it for demonstration; implement a cleanup strategy later if needed.
    # os.remove(file_path)

    return Response(content=content, media_type=media_type, headers={
        "Content-Disposition": f"attachment; filename={filename}"
    })


def save_document_and_get_download_link(document_obj: docx.Document, doc_type: str) -> str:
    """Saves a docx.Document object to a temporary file and returns its download URL."""
    # Ensure the directory for generated documents exists
    output_dir = "generated_documents"
    os.makedirs(output_dir, exist_ok=True)

    unique_filename = f"{doc_type}_{uuid.uuid4().hex}.docx"
    file_path = os.path.join(output_dir, unique_filename)
    document_obj.save(file_path)
    print(f"Document saved to: {file_path}")
    return f"http://127.0.0.1:8000/download/{unique_filename}"


# General Document Generation Endpoint (The fallback)
@app.post("/generate_document", response_model=GeneralResponse)
def generate_document_endpoint(request: DocumentRequest):
    """Generates a general document based on a complete DocumentRequest object."""
    print(f"Backend: Received a general document request for type: '{request.doc_type}'")
    try:
        output_document_obj = document_agent.generate_document(request)
        
        temp_file_path = "temp_generated_document.docx"
        output_document_obj.save(temp_file_path)
        
        doc = docx.Document(temp_file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        os.remove(temp_file_path)

        if not full_text:
            raise HTTPException(status_code=500, detail="Failed to extract text from generated document.")
        return GeneralResponse(result=full_text)
    except Exception as e:
        print(f"Error in generate_document_endpoint: {e}")
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")

# Dedicated Cover Letter Endpoint
@app.post("/create_cover_letter", response_model=GeneralResponse)
def create_cover_letter_endpoint(request: DocumentRequest):
    """Generates a cover letter from structured data."""
    print(f"Backend: Received a request to create a cover letter for: '{request.topic}'")
    try:
        request.doc_type = "cover_letter"
        output_document_obj = document_agent.generate_document(request)
        
        # Save the generated docx.Document object to a temp file and read its text
        temp_file_path = "temp_cover_letter.docx"
        output_document_obj.save(temp_file_path)
        
        doc = docx.Document(temp_file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        os.remove(temp_file_path)
        
        return GeneralResponse(result=full_text)
    except Exception as e:
        print(f"Error in create_cover_letter_endpoint: {e}")
        raise HTTPException(status_code=500, detail=f"Cover letter creation failed: {str(e)}")


# Dedicated Meeting Minutes Endpoint
@app.post("/create_minutes", response_model=GeneralResponse)
def create_minutes_endpoint(request: DocumentRequest):
    """Generates meeting minutes from structured data."""
    print(f"Backend: Received a request to create minutes for: '{request.topic}'")
    try:
        request.doc_type = "minutes"
        output_document_obj = document_agent.generate_document(request)
        
        # Save the generated docx.Document object to a temp file and read its text
        temp_file_path = "temp_minutes.docx"
        output_document_obj.save(temp_file_path)
        
        doc = docx.Document(temp_file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        os.remove(temp_file_path)
        
        return GeneralResponse(result=full_text)
    except Exception as e:
        print(f"Error in create_minutes_endpoint: {e}")
        raise HTTPException(status_code=500, detail=f"Meeting minutes creation failed: {str(e)}")

# Dedicated Memo Endpoint
@app.post("/create_memo", response_model=GeneralResponse)
def create_memo_endpoint(request: DocumentRequest):
    """Generates a memorandum from structured data."""
    print(f"Backend: Received a request to create a memo on topic: '{request.topic}'")
    try:
        request.doc_type = "memo"
        output_document_obj = document_agent.generate_document(request)
        
        # Save the generated docx.Document object to a temp file and read its text
        temp_file_path = "temp_memo.docx"
        output_document_obj.save(temp_file_path)
        
        doc = docx.Document(temp_file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        os.remove(temp_file_path)
        
        return GeneralResponse(result=full_text)
    except Exception as e:
        print(f"Error in create_memo_endpoint: {e}")
        raise HTTPException(status_code=500, detail=f"Memo creation failed: {str(e)}")
    
# Dedicated endpoint for Report Generation
@app.post("/create_report", response_model=GeneralResponse)
def create_report_endpoint(request: ProcessRequest):
    """Creates a report based on the provided prompt."""
    print("Backend: Received request to create a report.")
    try:
        # Assuming ReportAgent.create_report_content expects a 'topic'
        # The prompt from ProcessRequest is used as the topic.
        output_content = report_agent.create_report(
            topic=request.prompt, 
            tone="professional", # Default tone
            length="standard"    # Default length
        )
        if not output_content:
            raise HTTPException(status_code=500, detail="Failed to generate report content.")
        return GeneralResponse(result=output_content)
    except Exception as e:
        print(f"Error in create_report_endpoint: {e}")
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")
    
# In backend_server.py

# @app.post("/create_memo", response_model=GeneralResponse)
# def create_memo_endpoint(request: DocumentRequest):
#     """
#     Specifically generates a memorandum based on the provided structured data.
#     """
#     print(f"Backend: Received request to create a memo on topic: '{request.topic}'")
#     try:
#         # Use the document agent to generate the memo content
#         # Note: The agent handles the specifics of the 'memo' doc_type internally.
#         output_document_obj = document_agent.generate_document(request)
        
#         # Save the generated docx.Document object to a temp file and read its text
#         temp_file_path = "temp_generated_memo.docx"
#         output_document_obj.save(temp_file_path)
        
#         doc = docx.Document(temp_file_path)
#         full_text = "\n".join([para.text for para in doc.paragraphs])
#         os.remove(temp_file_path)

    #     if not full_text:
    #         raise HTTPException(status_code=500, detail="Failed to extract text from generated memo.")
            
    #     return GeneralResponse(result=full_text)
    # except Exception as e:
    #     print(f"Error in create_memo_endpoint: {e}")
    #     raise HTTPException(status_code=500, detail=f"Memo creation failed: {str(e)}")
    

# Dedicated endpoint for Data Analysis
@app.post("/analyze", response_model=GeneralResponse)
def analyze_endpoint(request: ProcessRequest):
    """Analyzes provided text content."""
    print("Backend: Received request to analyze content.")
    try:
        if not request.content:
            raise HTTPException(status_code=400, detail="No content provided for analysis.")
        
        # The data_agent.analyze_input_content generates the analysis text directly
        output_content = data_agent.analyze_input(raw_input=request.content, user_question=request.prompt)
        
        if not output_content:
            raise HTTPException(status_code=500, detail="Failed to generate analysis content.")
        return GeneralResponse(result=output_content)
    except Exception as e:
        print(f"Error in analyze_endpoint: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

# Dedicated endpoint for Summarization
@app.post("/summarize", response_model=GeneralResponse)
def summarize_endpoint(request: ProcessRequest):
    """Receives text content and returns a summary."""
    print("Backend: Received request to summarize document.")
    try:
        if not request.content:
            raise HTTPException(status_code=400, detail="No content provided for summarization.")
            
        # Corrected: Using generate_response instead of get_completion
        summary = llm_client.generate_response(
            f"Please provide a concise summary of the following document:\n\n{request.content}"
        )
        if not summary:
            raise HTTPException(status_code=500, detail="Failed to generate summary.")
        return GeneralResponse(result=summary)
    except Exception as e:
        print(f"Error during summarization: {e}")
        raise HTTPException(status_code=500, detail=f"Summarization failed: {str(e)}")

# General fallback endpoint for unstructured prompts
@app.post("/process", response_model=GeneralResponse)
def process_general_prompt(request: ProcessRequest):
    """Handles general prompts that don't fit other dedicated endpoints."""
    print(f"Backend: Received general prompt: '{request.prompt}'. Content present: {len(request.content) > 0}")
    try:
        # Corrected: Using generate_response instead of get_completion
        output_content = llm_client.generate_response(request.prompt)
        if not output_content:
            raise HTTPException(status_code=500, detail="Failed to get completion for general prompt.")
        return GeneralResponse(result=output_content)
    except Exception as e:
        print(f"Error in general process endpoint: {e}")
        raise HTTPException(status_code=500, detail=f"General prompt processing failed: {str(e)}")

# This block allows running the server directly for development
def main():
    """
    This is the main entry point for the bundled executable.
    It configures and runs the Uvicorn server.
    """
    import uvicorn
    
    # Check if running in a bundled environment
    is_bundled = getattr(sys, 'frozen', False)
    
    if is_bundled:
        print("Starting backend server from bundled executable...")
        # In a bundle, 'reload' must be False.
        # We can also add file logging for the production server.
        uvicorn.run(
            "wps_addin.backend_server:app",  # Point to the app object
            host="127.0.0.1",
            port=8000,
            reload=False,
            log_level="info"
        )
    else:
        # Development mode
        print("Starting backend server for development with auto-reload...")
        uvicorn.run(
            "wps_addin.backend_server:app",
            host="127.0.0.1",
            port=8000,
            reload=True,
            log_level="info"
        )

if __name__ == "__main__":
    main()
    
    
    # uvicorn wps_addin.backend_server:app --reload (use to run the server)

