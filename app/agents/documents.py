# import os
# from typing import List, Optional
# from docx import Document
# from docx.shared import Inches, Pt
# from pydantic import BaseModel, Field
# from app.llm_client import LLMClient 

# def generate_content(prompt: str) -> str:
#     """
#     Placeholder function to simulate AI content generation.
#     Replace this with your actual LLM client call.
#     """
#     print(f"Sending prompt to AI: \n{prompt}\n")
#     # Dummy response for demonstration purposes.
#     if "cover letter" in prompt:
#         return (
#             "I am writing to express my keen interest in the Senior Data Analyst position I saw advertised on LinkedIn. "
#             "With over five years of experience in data analytics and a proven track record of deriving actionable insights "
#             "from complex datasets, I am confident that I possess the skills and experience you are looking for."
#         )
#     elif "meeting minutes" in prompt:
#         return (
#             "Attendees: John Doe, Jane Smith, Peter Jones\n"
#             "Agenda Item 1: Q3 Financial Review - John Doe presented the quarterly earnings. Revenue is up by 15%.\n"
#             "Action Item: Jane Smith to prepare the detailed financial report by next Friday."
#         )
#     elif "memo" in prompt:
#         return (
#             "The purpose of this memo is to announce the new work-from-home policy, which will take effect on the first of next month. "
#             "All employees will be eligible to work remotely up to two days per week. Please coordinate with your department head to schedule your remote work days."
#         )
#     return "This is placeholder content based on the request."


# # Pydantic Model for Input Structure 
# class DocumentRequest(BaseModel):
#     """Structure for document generation requests."""
#     doc_type: str = Field(..., description="Type of document, e.g., 'cover_letter', 'minutes', 'memo'.")
#     topic: str = Field(..., description="The main subject or title of the document.")
#     length: str = Field('medium', description="Desired length: 'short', 'medium', or 'long'.")
#     tone: str = Field('formal', description="Tone of the content: 'formal', 'casual', 'technical'.")
#     audience: str = Field(..., description="The intended audience for the document, e.g., 'Hiring Manager', 'Project Team'.")
#     data_sources: Optional[List[str]] = Field(None, description="List of data sources or key points to include.")
#     template: Optional[str] = Field(None, description="Optional template identifier.")


# # The Document Generation Agent 
# class DocumentGenerationAgent:
#     """
#     An agent that generates Word documents (.docx) based on a structured request.
#     Uses LLM to generate the textual content for the document.
#     """
#     def __init__(self, llm_client: LLMClient):
#         self.llm_client = llm_client
#         print("Document Generation Agent initialized.")

#     def _create_cover_letter(self, document: Document, request: DocumentRequest):
#         """Generates content for a cover letter and adds it to the document."""
#         document.add_heading(request.topic, level=1) # Main heading for the letter subject

#         # More explicit prompt for LLM to generate cover letter body
#         prompt = (
#             f"Generate the full content of a {request.tone}, {request.length} cover letter. "
#             f"The purpose of the letter is an application for: '{request.topic}'. "
#             f"The letter is addressed to: '{request.audience}'. "
#             "Start with a formal salutation (e.g., 'Dear [Hiring Manager],'). "
#             "The body should explain keen interest, highlight relevant skills/experience, and express enthusiasm. "
#             "Conclude with a professional closing (e.g., 'Sincerely, [Your Name]'). "
#             "Output only the letter content, exactly as it should appear, including salutation and closing."
#         )
        
#         if request.data_sources:
#             prompt += f"\nKey qualifications/points to specifically include: {', '.join(request.data_sources)}."

#         content = self.get_llm_response(prompt) # Get actual content from LLM
#         document.add_paragraph(content)

#         print("Cover letter content generated and added.")

#     def _create_minutes(self, document: Document, request: DocumentRequest):
#         """Generates content for meeting minutes and adds it to the document."""
#         document.add_heading(f"Meeting Minutes: {request.topic}", level=1)
        
#         # Add a placeholder for date and attendees
#         document.add_paragraph(f"Date: [Current Date]")
#         document.add_paragraph(f"Attendees: {request.audience}")
#         document.add_paragraph(f"Meeting Topic: {request.topic}")
#         document.add_heading("Discussion & Actions", level=2)

#         prompt = (
#             f"Generate concise, structured meeting minutes for a meeting about '{request.topic}'. "
#             f"The tone should be {request.tone}. Attendees: {request.audience}. "
#             "Include key discussion points and clearly list action items with assigned persons/deadlines. "
#             "Format with bullet points for discussions and numbered lists for action items. "
#             "Output only the minutes content, starting directly with discussions."
#         )
#         if request.data_sources:
#             prompt += f"\nSpecific agenda items/points discussed: {', '.join(request.data_sources)}."
        
#         minutes_body = self._get_llm_response(prompt) # Get actual content from LLM
#         document.add_paragraph(minutes_body)
#         print("Meeting minutes content generated and added.")

#     def _create_memo(self, document: Document, request: DocumentRequest):
#         """Generates content for a memorandum and adds it to the document."""
#         # Set up standard memo header
#         document.add_heading("MEMORANDUM", level=0) # level 0 for main title
#         document.add_paragraph() # Spacer

#         table = document.add_table(rows=4, cols=2)
#         table.cell(0, 0).text = "TO:"
#         table.cell(0, 1).text = request.audience
#         table.cell(1, 0).text = "FROM:"
#         table.cell(1, 1).text = "[Your Name/Department]"
#         table.cell(2, 0).text = "DATE:"
#         table.cell(2, 1).text = "[Current Date]"
#         table.cell(3, 0).text = "SUBJECT:"
#         table.cell(3, 1).text = request.topic

#         # Add some spacing before body
#         document.add_paragraph()

#         prompt = (
#             f"Write the body of a {request.tone} memo. "
#             f"The memo is addressed to: {request.audience}. "
#             f"The subject is: '{request.topic}'. "
#             f"The desired length is {request.length}. "
#             "Start directly with the memo's main purpose and provide clear, concise information. Output only the memo body."
#         )
#         memo_body = self._get_llm_response(prompt) # Get actual content from LLM
#         document.add_paragraph(memo_body)
#         print("Memo content generated and added.")

#     def generate_document(self, request: DocumentRequest) -> Document:
#         """
#         Main method to generate a Word document based on the request.
#         """
#         document = Document()
#         # Set default document style and font
#         style = document.styles['Normal']
#         style.font.name = 'Calibri'
#         style.font.size = Pt(11)
        
#         # Set section margins if needed
#         # section = document.sections[0]
#         # section.left_margin = Inches(1)
#         # section.right_margin = Inches(1)
#         # section.top_margin = Inches(1)
#         # section.bottom_margin = Inches(1)

#         doc_type_map = {
#             'cover_letter': self._create_cover_letter,
#             'minutes': self._create_minutes,
#             'memo': self._create_memo,
#         }

#         generator_func = doc_type_map.get(request.doc_type.lower())

#         if generator_func:
#             generator_func(document, request)
#         else:
#             # Fallback for unsupported document types
#             document.add_heading(f"Document for: {request.topic}", level=1)
#             document.add_paragraph(f"This is a general document for {request.audience}. "
#                                     f"Document type '{request.doc_type}' is not specifically handled by an agent, "
#                                     f"so generic content will be generated.")
            
#             # Generate generic content using the primary instruction
#             prompt = (f"Generate a {request.length}, {request.tone} document about '{request.topic}' "
#                         f"for {request.audience}. Output only the main body content.")
#             generic_body = self._get_llm_response(prompt)
#             document.add_paragraph(generic_body)

#             print(f"Warning: Document type '{request.doc_type}' not specifically handled. Creating a generic document.")

#         return document

# # Example Usage 
# if __name__ == "__main__":
#     # Initialize the agent
#     agent = DocumentGenerationAgent()

#     # Define a request for a Cover Letter
#     cover_letter_request = DocumentRequest(
#         doc_type='cover_letter',
#         topic='Application for Senior Data Analyst Role',
#         audience='Hiring Manager',
#         tone='formal',
#         length='medium',
#         data_sources=['5+ years of experience', 'Proficiency in Python and SQL', 'Strong background in statistical analysis']
#     )

#     # Generate the document
#     generated_doc = agent.generate_document(cover_letter_request)
    
#     # Save the document
#     output_filename = "Generated_Cover_Letter.docx"
#     generated_doc.save(output_filename)

#     print(f"\n✅ Successfully generated document: '{output_filename}'")
#     print(f"   Saved at: {os.path.abspath(output_filename)}")

#     # --- Another Example: Generating a Memo ---
#     print("\n" + "="*50 + "\n")

#     memo_request = DocumentRequest(
#         doc_type='memo',
#         topic='New Work-From-Home Policy',
#         audience='All Employees',
#         tone='formal',
#         length='short'
#     )
    
#     memo_doc = agent.generate_document(memo_request)
#     memo_filename = "Generated_Memo.docx"
#     memo_doc.save(memo_filename)

#     print(f"\n✅ Successfully generated document: '{memo_filename}'")
#     print(f"   Saved at: {os.path.abspath(memo_filename)}")



import os
from typing import List, Optional
from docx import Document
from docx.shared import Pt
from pydantic import BaseModel, Field
from app.agents.llm_client import LLMClient

# Pydantic Model for Input Structure
class DocumentRequest(BaseModel):
    """Structure for document generation requests."""
    doc_type: str = Field(..., description="Type of document, e.g., 'cover_letter', 'minutes', 'memo'.")
    topic: str = Field(..., description="The main subject or title of the document.")
    length: str = Field('medium', description="Desired length: 'short', 'medium', or 'long'.")
    tone: str = Field('formal', description="Tone of the content: 'formal', 'casual', 'technical'.")
    audience: str = Field(..., description="The intended audience for the document, e.g., 'Hiring Manager', 'Project Team'.")
    data_sources: Optional[List[str]] = Field([], description="List of data sources or key points to include.")
    template: Optional[str] = Field(None, description="Optional template identifier.")


# The Document Generation Agent
class DocumentGenerationAgent:
    """
    An agent that generates Word documents (.docx) based on a structured request.
    Uses LLM to generate the textual content for the document.
    """
    def __init__(self, llm_client: LLMClient):
        self.llm_client = llm_client
        print("Document Generation Agent initialized.")

    def _create_cover_letter(self, document: Document, request: DocumentRequest):
        """Generates content for a cover letter and adds it to the document."""
        document.add_heading(request.topic, level=1)

        prompt = (
            f"Generate the full content of a {request.tone}, {request.length} cover letter. "
            f"The purpose of the letter is an application for: '{request.topic}'. "
            f"The letter is addressed to: '{request.audience}'. "
            "Start with a formal salutation (e.g., 'Dear [Hiring Manager],'). "
            "The body should explain keen interest, highlight relevant skills/experience, and express enthusiasm. "
            "Conclude with a professional closing (e.g., 'Sincerely, [Your Name]'). "
            "Output only the letter content, exactly as it should appear, including salutation and closing."
        )
        
        if request.data_sources:
            prompt += f"\nKey qualifications/points to specifically include: {', '.join(request.data_sources)}."

        #  Use the llm_client
        content = self.llm_client.generate_response(prompt)
        document.add_paragraph(content)
        print("Cover letter content generated and added.")

    def _create_minutes(self, document: Document, request: DocumentRequest):
        """Generates content for meeting minutes and adds it to the document."""
        document.add_heading(f"Meeting Minutes: {request.topic}", level=1)
        
        document.add_paragraph(f"Date: [Current Date]")
        document.add_paragraph(f"Attendees: {request.audience}")
        document.add_paragraph(f"Meeting Topic: {request.topic}")
        document.add_heading("Discussion & Actions", level=2)

        prompt = (
            f"Generate concise, structured meeting minutes for a meeting about '{request.topic}'. "
            f"The tone should be {request.tone}. Attendees: {request.audience}. "
            "Include key discussion points and clearly list action items with assigned persons/deadlines. "
            "Format with bullet points for discussions and numbered lists for action items. "
            "Output only the minutes content, starting directly with discussions."
        )
        if request.data_sources:
            prompt += f"\nSpecific agenda items/points discussed: {', '.join(request.data_sources)}."
        
        # CORRECTED: Use the llm_client
        minutes_body = self.llm_client.generate_response(prompt)
        document.add_paragraph(minutes_body)
        print("Meeting minutes content generated and added.")

    def _create_memo(self, document: Document, request: DocumentRequest):
        """Generates content for a memorandum and adds it to the document."""
        document.add_heading("MEMORANDUM", level=0)
        document.add_paragraph()

        table = document.add_table(rows=4, cols=2)
        table.cell(0, 0).text = "TO:"
        table.cell(0, 1).text = request.audience
        table.cell(1, 0).text = "FROM:"
        table.cell(1, 1).text = "[Your Name/Department]"
        table.cell(2, 0).text = "DATE:"
        table.cell(2, 1).text = "[Current Date]"
        table.cell(3, 0).text = "SUBJECT:"
        table.cell(3, 1).text = request.topic

        document.add_paragraph()

        prompt = (
            f"Write the body of a {request.tone} memo. "
            f"The memo is addressed to: {request.audience}. "
            f"The subject is: '{request.topic}'. "
            f"The desired length is {request.length}. "
            "Start directly with the memo's main purpose and provide clear, concise information. Output only the memo body."
        )
        # CORRECTED: Use the llm_client
        memo_body = self.llm_client.generate_response(prompt)
        document.add_paragraph(memo_body)
        print("Memo content generated and added.")

    def generate_document(self, request: DocumentRequest) -> Document:
        """Main method to generate a Word document based on the request."""
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        doc_type_map = {
            'cover_letter': self._create_cover_letter,
            'minutes': self._create_minutes,
            'memo': self._create_memo,
        }

        generator_func = doc_type_map.get(request.doc_type.lower())

        if generator_func:
            generator_func(document, request)
        else:
            document.add_heading(f"Document for: {request.topic}", level=1)
            document.add_paragraph(f"This is a general document for {request.audience}. "
                                    f"Document type '{request.doc_type}' is not specifically handled by an agent, "
                                    f"so generic content will be generated.")
            
            prompt = (f"Generate a {request.length}, {request.tone} document about '{request.topic}' "
                        f"for {request.audience}. Output only the main body content.")
            # CORRECTED: Use the llm_client
            generic_body = self.llm_client.generate_response(prompt)
            document.add_paragraph(generic_body)

            print(f"Warning: Document type '{request.doc_type}' not specifically handled. Creating a generic document.")

        return document