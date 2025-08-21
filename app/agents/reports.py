import os
import re
import json
from typing import Optional
from datetime import datetime
from pydantic import BaseModel, Field
# Imports for creating .docx files
from docx import Document
from docx.shared import Pt, Inches

from app.agents.llm_client import LLMClient

# Pydantic Model for a Structured Report 
class ReportOutput(BaseModel):
    """Defines the structured output for a generated report."""
    title: str = Field(..., description="The clear and professional title of the report.")
    executive_summary: str = Field(..., description="A concise summary of the report's key findings and conclusions.")
    main_content: str = Field(..., description="The main body of the report, formatted with paragraphs and potentially including its own subheadings.")
    conclusion: str = Field(..., description="A concluding summary that wraps up the report.")

#  Helper Functions 
def _clean_json_response(text: str) -> str:
    """Helper to strip markdown backticks or other text from a JSON string."""
    match = re.search(r'```(?:json)?\s*({.*})\s*```', text, re.DOTALL)
    if match:
        return match.group(1).strip()
    return text.strip()

#  Prompt to Request JSON 
def build_report_prompt(topic: str, tone: str = "professional", length: str = "standard") -> str:
    """Builds a detailed prompt that instructs the LLM to return a structured JSON object for a report."""
    return f"""
You are an expert business analyst and report writer, known for clarity, structure, and actionable insights.
Your task is to write a comprehensive report on the topic provided.
You must return your response as a single, valid JSON object, with no other text or markdown before or after it.

The JSON object must have the following exact structure:
{{
    "title": "A clear and professional title for the report",
    "executive_summary": "A concise summary of the report's key findings and conclusions.",
    "main_content": "The main body of the report. Use \\n\\n for new paragraphs and start major sections with headings like '## Introduction' or '## Analysis'.",
    "conclusion": "A concluding summary that wraps up the report and may suggest next steps."
}}

---
TOPIC: "{topic}"
TONE: "{tone}"
DESIRED LENGTH: "{length}" (e.g., 'short' for a brief memo, 'standard' for 1-2 pages, 'comprehensive' for a detailed analysis)
---

Now, generate the final JSON response based on the instructions above.
"""

# Function to Save the Report to a .docx file 
def save_report_to_docx(report: ReportOutput, output_dir: str = "generated_reports") -> str:
    """
    Saves the generated report to a .docx file with professional formatting.

    Args:
        report (ReportOutput): The structured report data.
        output_dir (str): The directory to save the file in.

    Returns:
        str: The full path to the saved .docx file.
    """
    try:
        os.makedirs(output_dir, exist_ok=True)
        sanitized_title = re.sub(r'[\\/*?:"<>|]', "", report.title)[:50]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Report_{sanitized_title}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)

        document = Document()
        style = document.styles['Normal'].font
        style.name = 'Calibri'
        style.size = Pt(11)

        # Add Title
        document.add_heading(report.title, level=1)
        document.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d')}")

        # Add Executive Summary
        document.add_heading("Executive Summary", level=2)
        document.add_paragraph(report.executive_summary)

        # Add Main Content
        document.add_heading("Main Report", level=2)
        # Split content to preserve paragraphs and handle subheadings
        for paragraph_text in report.main_content.strip().split('\n\n'):
            stripped_para = paragraph_text.strip()
            if stripped_para.startswith('## '):
                # Add as a smaller heading if AI uses markdown-style headings
                document.add_heading(stripped_para.replace('## ', ''), level=3)
            elif stripped_para:
                document.add_paragraph(stripped_para)

        # Add Conclusion
        document.add_heading("Conclusion", level=2)
        document.add_paragraph(report.conclusion)

        document.save(filepath)
        print(f"-> Report successfully saved to: {filepath}")
        return filepath

    except Exception as e:
        print(f"Error saving document: {e}")
        raise

# ReportAgent Class 
class ReportAgent:
    def __init__(self, llm_client: LLMClient):
        """Initializes the agent with a pre-configured LLMClient."""
        self.llm_client = llm_client
        print("-> Report Agent activated.")

    def create_report(self, topic: str, tone: Optional[str] = "professional", length: Optional[str] = "1-2 pages") -> str:
        """
        Generates a report, saves it to a .docx file, and returns the filepath.
        """
        if not topic:
            raise ValueError("A topic must be provided to generate a report.")
            
        prompt = build_report_prompt(topic, tone, length)
        print(f"-> Generating report on '{topic}'...")
        raw_response = self.llm_client.generate_response(prompt, json_mode=True)
        cleaned_response = _clean_json_response(raw_response)
        
        try:
            parsed_json = json.loads(cleaned_response)
            report_output = ReportOutput.model_validate(parsed_json)
        except (json.JSONDecodeError, TypeError) as e:
            print(f"Error: Failed to parse the LLM's response. Saving an error report.")
            report_output = ReportOutput(
                title=f"Error Generating Report on '{topic}'",
                executive_summary="The AI's response could not be parsed correctly. This document contains the error details.",
                main_content=f"Error Details: {e}\n\nRaw AI Response:\n{raw_response}",
                conclusion="No conclusion could be generated due to the parsing error."
            )
        
        # Save the result (either the report or an error report) to a .docx file
        filepath = save_report_to_docx(report_output)
        return filepath
