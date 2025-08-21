"""
This is the agent that is responsible with data analysis and writing a report on the same.

"""
import io
import json
import re
import argparse
import os
from typing import Any, Dict, Optional, List
import pandas as pd
import numpy as np
import scipy.stats as stats # For p-values, etc.
import matplotlib.pyplot as plt
import seaborn as sns
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn # For font setting

from app.agents.llm_client import LLMClient 

# Pydantic Schemas for Structured Output 
class CorrelationResult(BaseModel):
    """Details for a single correlation."""
    variable1: str
    variable2: str
    correlation: float
    p_value: Optional[float] = None
    interpretation: str

class TTestResult(BaseModel):
    """Details for a single independent samples t-test."""
    group_column: str
    numeric_column: str
    group1_name: Any
    group2_name: Any
    t_statistic: float
    p_value: float
    interpretation: str

class AnovaResult(BaseModel):
    """Details for a single one-way ANOVA test."""
    group_column: str
    numeric_column: str
    f_statistic: float
    p_value: float
    interpretation: str
    group_means: Optional[Dict[Any, float]] = None # Means for each group

class ZTestResult(BaseModel):
    """Details for a single one-sample Z-test."""
    numeric_column: str
    hypothesized_mean: float
    z_statistic: float
    p_value: float
    interpretation: str

# Add this new class
class VisualizationRecommendation(BaseModel):
    """Details for a single recommended visualization."""
    chart_type: str = Field(..., description="The type of chart, e.g., 'boxplot', 'bar_chart', 'scatter_plot'.")
    columns: List[str] = Field(..., description="The columns from the data to be used for the chart.")
    description: str = Field(..., description="A brief explanation of what this chart would show and why it's useful.")


class StatisticalSummary(BaseModel):
    """Structured summary of statistical findings."""
    descriptive_stats: Dict[str, Any] = Field(..., description="Descriptive statistics (mean, median, std, etc.) per column.")
    correlations: List[CorrelationResult] = Field(..., description="Key correlations identified between numerical variables.")
    t_tests: List[TTestResult] = Field([], description="Results of independent samples t-tests.")
    anova_results: List[AnovaResult] = Field([], description="Results of one-way ANOVA tests.")
    z_tests: List[ZTestResult] = Field([], description="Results of one-sample Z-tests against a hypothesized mean.")


class AnalysisOutput(BaseModel):
    summary: str = Field(..., description="Short, human-readable summary of the dataset and main findings.")
    insights: list[str] = Field(..., description="List of concrete, data-driven insights discovered.")
    recommended_visualizations: List[VisualizationRecommendation] = Field(..., description="Suggested visualizations, specifying type and columns.")
    risk_flags: list[str] = Field(..., description="Potential risks, biases, or data quality issues to be aware of.")
    pandas_code_snippet: str = Field(..., description="A short, functional pandas code snippet to reproduce a key insight or chart.")
    # Fields to store the raw statistical results and plot paths for the report
    statistical_results: Optional[StatisticalSummary] = Field(None, description="Detailed statistical analysis results.")
    plot_image_paths: List[str] = Field([], description="List of file paths to generated visualization images.")


# Helper Functions for Data Handling
def try_parse_csv_or_table(text: str) -> Optional[pd.DataFrame]:
    """Tries to parse incoming text as CSV or TSV into a pandas DataFrame."""
    if not text or len(text.strip()) == 0: return None
    try:
        sample_lines = text.strip().splitlines()
        # Try to infer delimiter from first few lines
        sep = ',' if sample_lines[0].count(',') >= sample_lines[0].count('\t') else '\t'
        
        # Read with pandas, trying multiple engines for robustness
        try:
            df = pd.read_csv(io.StringIO(text), sep=sep, engine='c', on_bad_lines='skip')
        except Exception:
            df = pd.read_csv(io.StringIO(text), sep=sep, engine='python', on_bad_lines='skip')
            
        # Drop unnamed columns that might result from malformed CSVs
        df = df.loc[:, ~df.columns.str.contains('^Unnamed')]

        return df if not df.empty and not df.isnull().all().all() else None
    except Exception as e:
        print(f"Error parsing CSV/Table: {e}")
        return None

def get_local_data_summary(df: pd.DataFrame, max_rows: int = 5) -> Dict[str, Any]:
    """Computes deterministic summaries of a DataFrame to include in the model prompt."""
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    summary = {
        "row_count": int(len(df)),
        "column_count": int(df.shape[1]),
        "column_names": df.columns.tolist(),
        "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()},
        "first_rows_sample": df.head(max_rows).to_dict(orient="records"),
        "numeric_column_summary": df[numeric_cols].describe().to_dict() if numeric_cols else {},
        "missing_values_per_column": df.isnull().sum().to_dict(),
    }
    return summary

def _perform_statistical_analysis(df: pd.DataFrame) -> StatisticalSummary:
    """Performs core statistical analysis (descriptive, correlations, t-tests, ANOVA, Z-tests) on the DataFrame."""
    
    # Initialize containers for results
    descriptive_stats_dict = {}
    correlations_list: List[CorrelationResult] = []
    t_tests_list: List[TTestResult] = []
    anova_results_list: List[AnovaResult] = []
    z_tests_list: List[ZTestResult] = []

    numeric_df = df.select_dtypes(include=np.number)
    categorical_df = df.select_dtypes(include='object').copy() # Use copy to avoid SettingWithCopyWarning

    # Descriptive Statistics ---
    if not numeric_df.empty:
        descriptive_stats_dict = {
            'numeric_columns': numeric_df.columns.tolist(),
            'statistics': numeric_df.describe(include='all').to_dict()
        }
    
    # Correlations and P-values
    if not numeric_df.empty and numeric_df.shape[1] >= 2:
        corr_matrix = numeric_df.corr(method='pearson')
        for i in range(corr_matrix.shape[0]):
            for j in range(i + 1, corr_matrix.shape[1]):
                var1 = corr_matrix.columns[i]
                var2 = corr_matrix.columns[j]
                corr_val = corr_matrix.iloc[i, j]
                
                if not pd.isna(corr_val):
                    try:
                        # Drop NaNs for pearson's calculation
                        cleaned_data = numeric_df[[var1, var2]].dropna()
                        if len(cleaned_data) < 2: # Need at least 2 observations for correlation
                            continue

                        if cleaned_data[var1].std() == 0 or cleaned_data[var2].std() == 0:
                            p_value = np.nan # No variance, p-value undefined
                        else:
                            r, p_val = stats.pearsonr(cleaned_data[var1], cleaned_data[var2])
                            p_value = p_val
                        
                        interpretation = ""
                        abs_corr = abs(corr_val)
                        if abs_corr >= 0.7: interpretation = "Very Strong"
                        elif abs_corr >= 0.5: interpretation = "Strong"
                        elif abs_corr >= 0.3: interpretation = "Moderate"
                        elif abs_corr >= 0.1: interpretation = "Weak"
                        else: interpretation = "Very Weak/No"
                        
                        direction = "positive" if corr_val > 0 else "negative" if corr_val < 0 else ""
                        interpretation = f"{interpretation} {direction} correlation."
                        if p_value is not None and p_value < 0.05:
                            interpretation += " (Statistically significant at p < 0.05)"
                        elif p_value is not None and p_value >= 0.05:
                            interpretation += " (Not statistically significant at p < 0.05)"
                        
                        correlations_list.append(
                            CorrelationResult(
                                variable1=str(var1),
                                variable2=str(var2),
                                correlation=float(corr_val),
                                p_value=float(p_value) if p_value is not None else None,
                                interpretation=interpretation
                            )
                        )
                    except Exception as e:
                        print(f"Warning: Could not compute correlation/p-value for {var1}-{var2}: {e}")

    # T-Tests (Independent Samples) 
    # Look for numeric columns and categorical columns with exactly two unique values
    for num_col in numeric_df.columns:
        for cat_col in categorical_df.columns:
            unique_values = categorical_df[cat_col].dropna().unique()
            if len(unique_values) == 2:
                group1_name, group2_name = unique_values[0], unique_values[1]
                group1_data = df[df[cat_col] == group1_name][num_col].dropna()
                group2_data = df[df[cat_col] == group2_name][num_col].dropna()

                if len(group1_data) > 1 and len(group2_data) > 1: # Need at least 2 data points per group
                    try:
                        t_stat, p_val = stats.ttest_ind(group1_data, group2_data, equal_var=True) # Assume equal variance
                        
                        interpretation = f"Mean of '{num_col}' in '{group1_name}' vs '{group2_name}' groups."
                        if p_val < 0.05:
                            interpretation += f" There is a statistically significant difference (p < 0.05) between the means."
                        else:
                            interpretation += f" No statistically significant difference (p >= 0.05) found between the means."
                        
                        t_tests_list.append(
                            TTestResult(
                                group_column=str(cat_col),
                                numeric_column=str(num_col),
                                group1_name=str(group1_name),
                                group2_name=str(group2_name),
                                t_statistic=float(t_stat),
                                p_value=float(p_val),
                                interpretation=interpretation
                            )
                        )
                    except Exception as e:
                        print(f"Warning: Could not perform t-test for {num_col} by {cat_col}: {e}")

    # ANOVA (One-Way) 
    # Look for numeric columns and categorical columns with 3 or more (but not too many) unique values
    for num_col in numeric_df.columns:
        for cat_col in categorical_df.columns:
            unique_values = categorical_df[cat_col].dropna().unique()
            # Max 10 categories for readability and computational reasons
            if 2 < len(unique_values) <= 10: 
                groups = [df[df[cat_col] == g][num_col].dropna() for g in unique_values]
                groups = [g for g in groups if len(g) > 1] # Ensure groups have enough data

                if len(groups) == len(unique_values) and len(groups) > 1: # Ensure all original groups have data and there's more than one group
                    try:
                        f_stat, p_val = stats.f_oneway(*groups)
                        
                        interpretation = f"Comparing means of '{num_col}' across different groups in '{cat_col}'."
                        if p_val < 0.05:
                            interpretation += f" There is a statistically significant difference (p < 0.05) between at least two group means."
                        else:
                            interpretation += f" No statistically significant difference (p >= 0.05) found between group means."
                        
                        group_means = {str(g_name): float(df[df[cat_col] == g_name][num_col].mean()) for g_name in unique_values}

                        anova_results_list.append(
                            AnovaResult(
                                group_column=str(cat_col),
                                numeric_column=str(num_col),
                                f_statistic=float(f_stat),
                                p_value=float(p_val),
                                interpretation=interpretation,
                                group_means=group_means
                            )
                        )
                    except Exception as e:
                        print(f"Warning: Could not perform ANOVA for {num_col} by {cat_col}: {e}")

    # One-Sample Z-Tests
    # For each numeric column, perform a one-sample Z-test against a hypothesized mean of 0
    # Assuming sample standard deviation as estimate for population SD given large enough N
    for num_col in numeric_df.columns:
        data = numeric_df[num_col].dropna()
        if len(data) >= 30 and data.std() > 0: # Z-test typically for N > 30
            hypothesized_mean = 0.0 # Common null hypothesis: mean is zero
            sample_mean = data.mean()
            sample_std = data.std()
            n = len(data)

            try:
                z_statistic = (sample_mean - hypothesized_mean) / (sample_std / np.sqrt(n))
                p_value = stats.norm.sf(abs(z_statistic)) * 2 # Two-tailed p-value

                interpretation = f"Testing if the mean of '{num_col}' is significantly different from {hypothesized_mean}."
                if p_value < 0.05:
                    interpretation += f" Mean is statistically significantly different from {hypothesized_mean} (p < 0.05)."
                else:
                    interpretation += f" Mean is not statistically significantly different from {hypothesized_mean} (p >= 0.05)."
                
                z_tests_list.append(
                    ZTestResult(
                        numeric_column=str(num_col),
                        hypothesized_mean=float(hypothesized_mean),
                        z_statistic=float(z_statistic),
                        p_value=float(p_value),
                        interpretation=interpretation
                    )
                )
            except Exception as e:
                print(f"Warning: Could not perform Z-test for {num_col}: {e}")

    return StatisticalSummary(
        descriptive_stats=descriptive_stats_dict,
        correlations=correlations_list,
        t_tests=t_tests_list,
        anova_results=anova_results_list,
        z_tests=z_tests_list
    )


def _generate_plots(df: pd.DataFrame, output_dir: str) -> List[str]:
    """Generates a few common plots and saves them as PNG images."""
    os.makedirs(output_dir, exist_ok=True)
    plot_paths = []
    
    # Set a consistent style for plots
    sns.set_style("whitegrid")
    plt.rcParams["figure.figsize"] = (8, 6) # Default figure size

    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    categorical_cols = df.select_dtypes(include='object').columns.tolist()
    
    # Plot 1: Histograms for numeric columns
    for col in numeric_cols:
        if df[col].nunique() > 1: # Only plot if there's variance
            plt.figure()
            sns.histplot(df[col].dropna(), kde=True)
            plt.title(f'Distribution of {col}')
            plt.xlabel(col)
            plt.ylabel('Frequency')
            plot_path = os.path.join(output_dir, f'hist_{col}.png')
            plt.savefig(plot_path, bbox_inches='tight')
            plt.close()
            plot_paths.append(plot_path)

    # Plot 2: Pairplot for a few numeric columns (if many, select a subset)
    if len(numeric_cols) >= 2:
        subset_cols = numeric_cols[:min(len(numeric_cols), 5)] # Limit to max 5 for pairplot performance/readability
        if len(subset_cols) >= 2:
            try:
                # Ensure no non-numeric data creeps in, though dtypes should handle this
                valid_pair_df = df[subset_cols].dropna() 
                if not valid_pair_df.empty:
                    pairplot_fig = sns.pairplot(valid_pair_df)
                    plt.suptitle('Pairplot of Key Numeric Variables', y=1.02) # Adjust title position
                    plot_path = os.path.join(output_dir, 'pairplot.png')
                    plt.savefig(plot_path, bbox_inches='tight')
                    plt.close()
                    plot_paths.append(plot_path)
            except Exception as e:
                print(f"Warning: Could not generate pairplot: {e}")

    # Plot 3: Box plots for numeric by categorical (if applicable)
    if numeric_cols and categorical_cols:
        for num_col in numeric_cols[:min(len(numeric_cols), 3)]: # Take a few numeric cols
            for cat_col in categorical_cols[:min(len(categorical_cols), 2)]: # Take a few categorical cols
                if df[cat_col].nunique() < 10 and df[num_col].nunique() > 1: # Limit for readability
                    plt.figure()
                    sns.boxplot(x=cat_col, y=num_col, data=df)
                    plt.title(f'{num_col} by {cat_col}')
                    plt.xlabel(cat_col)
                    plt.ylabel(num_col)
                    plt.xticks(rotation=45, ha='right') # Rotate labels for better fit
                    plot_path = os.path.join(output_dir, f'boxplot_{num_col}_by_{cat_col}.png')
                    plt.savefig(plot_path, bbox_inches='tight')
                    plt.close()
                    plot_paths.append(plot_path)

    # Plot 4: Count plots for categorical columns
    for col in categorical_cols:
        if df[col].nunique() < 15: # Limit for readability
            plt.figure()
            sns.countplot(y=col, data=df, order=df[col].value_counts().index) # Order by frequency
            plt.title(f'Count of {col}')
            plt.xlabel('Count')
            plt.ylabel(col)
            plot_path = os.path.join(output_dir, f'countplot_{col}.png')
            plt.savefig(plot_path, bbox_inches='tight')
            plt.close()
            plot_paths.append(plot_path)

    return plot_paths

def _clean_json_response(text: str) -> str:
    """Helper to strip markdown backticks or other text from a JSON string."""
    match = re.search(r'```(?:json)?\s*({.*})\s*```', text, re.DOTALL)
    if match:
        return match.group(1).strip()
    
    match = re.search(r'({.*})', text.strip(), re.DOTALL)
    if match:
        return match.group(1).strip()

    return text.strip() 

def _create_analysis_report_docx(
    df: pd.DataFrame, # Original DataFrame for context
    analysis_output: AnalysisOutput, 
    output_filepath: str,
    output_dir_for_plots: str # Need path to plots
) -> str:
    """
    Creates a Word document (.docx) containing the analysis report.
    Embeds LLM insights, statistical findings, and generated plots.
    """
    document = Document()

    # Set default font and size
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    # Set page orientation to landscape for wider plots if preferred
    section = document.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.orientation = WD_ORIENT.LANDSCAPE

    # Title Page 
    document.add_heading('Data Analysis Report', level=0)
    document.add_paragraph('Generated by AI Office Automation Assistant')
    document.add_paragraph(f'Date: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
    document.add_page_break()

    # Executive Summary 
    document.add_heading('1. Executive Summary', level=1)
    document.add_paragraph(analysis_output.summary)
    document.add_page_break()

    # Key Insights 
    document.add_heading('2. Key Insights', level=1)
    for i, insight in enumerate(analysis_output.insights, 1):
        document.add_paragraph(f'• {insight}')
    document.add_page_break()

    # Statistical Findings 
    document.add_heading('3. Statistical Findings', level=1)
    if analysis_output.statistical_results:
        # Descriptive Statistics
        document.add_heading('3.1. Descriptive Statistics', level=2)
        
        if analysis_output.statistical_results.descriptive_stats.get('statistics'): # Check for 'statistics' key
            document.add_paragraph("Summary of Numeric Columns:")
            # Convert the describe() output dict to DataFrame for easier display in docx
            desc_df = pd.DataFrame(analysis_output.statistical_results.descriptive_stats['statistics'])
            
            # Add table for descriptive stats
            table = document.add_table(rows=desc_df.shape[0] + 1, cols=desc_df.shape[1] + 1)
            table.style = 'Table Grid' # Apply a basic table style

            # Add header row
            table.cell(0, 0).text = "Statistic"
            for col_idx, col_name in enumerate(desc_df.columns):
                table.cell(0, col_idx + 1).text = str(col_name)

            # Add data rows
            for row_idx, (index_name, row_data) in enumerate(desc_df.iterrows()):
                table.cell(row_idx + 1, 0).text = str(index_name)
                for col_idx, value in enumerate(row_data):
                    table.cell(row_idx + 1, col_idx + 1).text = f"{value:.2f}" if isinstance(value, (int, float)) else str(value)
            
            document.add_paragraph() # Add space after table


        # Correlations
        document.add_heading('3.2. Correlations', level=2)
        if analysis_output.statistical_results.correlations:
            for corr in analysis_output.statistical_results.correlations:
                p = document.add_paragraph()
                p.add_run(f'{corr.variable1} vs {corr.variable2}: ').bold = True
                p.add_run(f'r={corr.correlation:.3f}')
                if corr.p_value is not None and not np.isnan(corr.p_value): # Check for NaN p_value
                    p.add_run(f', p={corr.p_value:.3f}')
                p.add_run(f' ({corr.interpretation})')
        else:
            document.add_paragraph("No significant correlations found between numeric variables or insufficient numeric data.")
    
        #  T-Tests 
        document.add_heading('3.3. T-Tests', level=2)
        if analysis_output.statistical_results.t_tests:
            for tt in analysis_output.statistical_results.t_tests:
                p = document.add_paragraph()
                p.add_run(f"Comparison of '{tt.numeric_column}' between '{tt.group1_name}' and '{tt.group2_name}' in '{tt.group_column}': ").bold = True
                p.add_run(f"t-statistic={tt.t_statistic:.3f}, p-value={tt.p_value:.3f}. ")
                p.add_run(f"Interpretation: {tt.interpretation}")
        else:
            document.add_paragraph("No relevant t-tests performed or statistically significant differences found.")

        # ANOVA Results 
        document.add_heading('3.4. ANOVA Results', level=2)
        if analysis_output.statistical_results.anova_results:
            for anova in analysis_output.statistical_results.anova_results:
                p = document.add_paragraph()
                p.add_run(f"ANOVA for '{anova.numeric_column}' grouped by '{anova.group_column}': ").bold = True
                p.add_run(f"F-statistic={anova.f_statistic:.3f}, p-value={anova.p_value:.3f}. ")
                p.add_run(f"Interpretation: {anova.interpretation}")
                if anova.group_means:
                    means_str = ", ".join([f"{k}: {v:.2f}" for k,v in anova.group_means.items()])
                    p.add_run(f" Group Means: {{{means_str}}}.")
        else:
            document.add_paragraph("No relevant ANOVA tests performed or statistically significant differences found across groups.")

        # Z-Tests 
        document.add_heading('3.5. Z-Tests', level=2)
        if analysis_output.statistical_results.z_tests:
            for zt in analysis_output.statistical_results.z_tests:
                p = document.add_paragraph()
                p.add_run(f"One-Sample Z-Test for '{zt.numeric_column}' (Hypothesized Mean={zt.hypothesized_mean}): ").bold = True
                p.add_run(f"Z-statistic={zt.z_statistic:.3f}, p-value={zt.p_value:.3f}. ")
                p.add_run(f"Interpretation: {zt.interpretation}")
        else:
            document.add_paragraph("No relevant Z-tests performed or sufficient data for Z-tests.")


    else: # If analysis_output.statistical_results is None
        document.add_paragraph("No detailed statistical analysis performed or available.")
    
    document.add_page_break()


    # --- Visualizations ---
    document.add_heading('4. Visualizations', level=1)
    if analysis_output.plot_image_paths:
        for plot_path in analysis_output.plot_image_paths:
            if os.path.exists(plot_path):
                try:
                    # Add plot title based on filename
                    plot_title = os.path.basename(plot_path).replace('.png', '').replace('_', ' ').title()
                    p = document.add_paragraph()
                    r = p.add_run(plot_title)
                    r.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    document.add_picture(plot_path, width=Inches(6.5)) # Adjust width as needed
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    document.add_paragraph() # Add space after image
                except Exception as e:
                    document.add_paragraph(f"Could not embed image {os.path.basename(plot_path)}: {e}")
            else:
                document.add_paragraph(f"Warning: Image file not found: {os.path.basename(plot_path)}")
    else:
        document.add_paragraph("No visualizations were generated for this dataset.")
    document.add_page_break()

    # Recommended Visualizations (from LLM)
    document.add_heading('5. Recommended Visualizations', level=1)
    if analysis_output.recommended_visualizations:
        for i, rec_viz in enumerate(analysis_output.recommended_visualizations, 1):
            document.add_paragraph(f'• {rec_viz}')
    else:
        document.add_paragraph("No specific visualizations recommended by the AI.")
    document.add_page_break()


    # Risk Flags
    document.add_heading('6. Risk Flags', level=1)
    if analysis_output.risk_flags:
        for i, flag in enumerate(analysis_output.risk_flags, 1):
            document.add_paragraph(f'• {flag}')
    else:
        document.add_paragraph("No specific risk flags identified by the AI.")
    document.add_page_break()

    # --- Pandas Code Snippet ---
    document.add_heading('7. Pandas Code Snippet', level=1)
    if analysis_output.pandas_code_snippet:
        document.add_paragraph("Here's a relevant Python Pandas code snippet:")
        code_paragraph = document.add_paragraph()
        code_paragraph.add_run(analysis_output.pandas_code_snippet)
        # You might want to apply a specific style for code if defined in your docx template
        code_paragraph.style = 'Intense Quote' # A built-in style that can look like code
    else:
        document.add_paragraph("No specific Pandas code snippet provided.")
    document.add_page_break()

    try:
        document.save(output_filepath)
        print(f"Analysis report saved to: {output_filepath}")
        return output_filepath
    except Exception as e:
        print(f"Error saving analysis report: {e}")
        raise


# Main Agent Class 
class StructuredDataAgent:
    def __init__(self, llm_client: LLMClient):
        """
        Initializes the agent with a pre-configured LLMClient.

        Args:
            llm_client (LLMClient): An instance of the LLMClient.
        """
        self.llm_client = llm_client
        # Directory to save plots temporarily
        self.temp_plot_dir = "temp_plots"
        os.makedirs(self.temp_plot_dir, exist_ok=True)
        print(f"Initialized StructuredDataAgent using provider: {self.llm_client.provider}, model: {self.llm_client.model}")


    def analyze_input(self, raw_input: str, user_question: str = "") -> str: # Now returns path to docx
        """Main entry point for analyzing tabular data, generating visualizations, and creating a report."""
        df = try_parse_csv_or_table(raw_input)
        if df is None:
            raise ValueError("Input could not be parsed as a valid CSV or table.")
        
        # Perform Statistical Analysis
        statistical_results = _perform_statistical_analysis(df)
        
        # Generate Plots
        plot_image_paths = _generate_plots(df, self.temp_plot_dir)

        # Build Prompt for LLM with all available information
        data_summary = get_local_data_summary(df)
        prompt = self._build_llm_analysis_prompt(data_summary, statistical_results, user_question)

        # Get LLM's structured analysis (summary, insights, etc.)
        llm_raw_response_for_analysis = self.llm_client.generate_response(prompt=prompt, json_mode=True) 
        cleaned_response_for_analysis = _clean_json_response(llm_raw_response_for_analysis)
        
        try:
            parsed_analysis_json = json.loads(cleaned_response_for_analysis)
            # Ensure the statistical_results part is included for Pydantic validation
            parsed_analysis_json['statistical_results'] = statistical_results.dict()
            parsed_analysis_json['plot_image_paths'] = plot_image_paths # Also explicitly ensure paths are there

            llm_analysis_output = AnalysisOutput.model_validate(parsed_analysis_json)
        except (json.JSONDecodeError, TypeError) as e:
            print(f"Error parsing LLM analysis response: {e}")
            print(f"Raw LLM response: {cleaned_response_for_analysis}")
            raise RuntimeError(f"Failed to get structured analysis from LLM: {e}")

        # 5. Create Word Document Report
        output_filepath = os.path.join(os.getcwd(), f"Analysis_Report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx")
        generated_report_path = _create_analysis_report_docx(df, llm_analysis_output, output_filepath, self.temp_plot_dir)

        # Clean up temporary plot images
        for p in plot_image_paths:
            if os.path.exists(p):
                os.remove(p)
        if os.path.exists(self.temp_plot_dir) and not os.listdir(self.temp_plot_dir):
            os.rmdir(self.temp_plot_dir)

        return generated_report_path

    def _build_llm_analysis_prompt(self, data_summary: Dict[str, Any], statistical_results: StatisticalSummary, user_question: str = "") -> str:
        """
        Builds a comprehensive prompt for the LLM to generate the structured analysis.
        """
        prompt = (
            "You are an expert data analyst. Based on the following dataset summary and statistical findings, "
            "provide a structured JSON response following the AnalysisOutput schema. "
            "Do NOT include any text or markdown formatting outside the JSON object itself. "
            "Ensure ALL keys from the AnalysisOutput schema are present and correctly typed (even if empty lists)."
            "\n\nDATASET OVERVIEW (for your reference):\n"
            f"{json.dumps(data_summary, indent=2)}\n\n"
            "STATISTICAL FINDINGS (interpret these in your analysis):\n"
            f"{json.dumps(statistical_results.dict(), indent=2)}\n\n"
            "USER REQUEST:\n"
            f"{user_question or 'Perform a thorough analysis. Identify key trends, correlations, potential issues, and suggest relevant visualizations.'}\n\n"
            "Your response must be a single JSON object with the following keys: "
            "'summary' (string), 'insights' (list of strings), 'recommended_visualizations' (list of objects), "
            "'risk_flags' (list of strings), and 'pandas_code_snippet' (string)."
            "\n\nRULES FOR SPECIFIC KEYS:"
            "\n- 'insights': Must be a JSON list of individual, complete sentences as strings. For example: [\"Insight one.\", \"Insight two.\"]"
            "\n- 'recommended_visualizations': Must be a JSON list of objects. Each object must have three keys: "
            "'chart_type' (string), 'columns' (list of strings), and 'description' (string)."
        )
        return prompt


# Command-Line Interface for Standalone Testing 
if __name__ == "__main__":
    # Ensure all required libraries are installed:
    # pip install pandas numpy scipy matplotlib seaborn python-docx pydantic python-dotenv requests google-generativeai openai

    parser = argparse.ArgumentParser(description="Structured Data Analysis Agent CLI")
    parser.add_argument("--file", "-f", required=True, help="Path to the input CSV file.")
    parser.add_argument("--question", "-q", default="Perform a thorough data analysis.", help="Your question for the analysis.")
    parser.add_argument("--provider", "-p", default="gemini", choices=["ollama", "openai", "gemini"], help="LLM provider to use.")
    
    args = parser.parse_args()

    try:
        print(f"Loading data from: {args.file}")
        with open(args.file, "r", encoding="utf-8") as f:
            raw_data = f.read()

        # Instantiate LLM client for testing
        client = LLMClient(provider=args.provider)
        
        # Pass client to the agent
        agent = StructuredDataAgent(llm_client=client)
        
        print(f"Analyzing data with: {args.provider} model: {client.model}")
        
        # The analyze_input now returns the path to the DOCX report
        generated_report_path = agent.analyze_input(raw_data, user_question=args.question)
        
        print(f"\n✅ Analysis Report Generated and Saved to: {generated_report_path}")

        # Optional: Prompt to open the file
        open_file = input("Do you want to open the generated report now? (y/n): ").lower()
        if open_file == 'y':
            try:
                import subprocess
                if sys.platform == "win32":
                    os.startfile(generated_report_path)
                elif sys.platform == "darwin": # macOS
                    subprocess.call(('open', generated_report_path))
                elif sys.platform.startswith('linux'):
                    subprocess.call(('xdg-open', generated_report_path))
                else:
                    print("Could not automatically open file. Please open it manually.")
            except Exception as e:
                print(f"Error opening file: {e}")

    except (ValueError, RuntimeError, FileNotFoundError) as e:
        print(f"\n❌ An error occurred while analyzing your data: {e}")
